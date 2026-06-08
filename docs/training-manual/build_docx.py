"""Typeset the Practitioner's Manual markdown into a styled .docx.

Pure-Python (no pandoc): python-docx for the document, svglib+reportlab to rasterize the
four SVG figures to PNG. Assembles the six part files into one clean manuscript, dropping the
reviewer scaffolding (styling-note comments, per-part metadata blocks, "About this part" and
"End of Part — what to review" sections).

Run with the build deps overlaid:
    uv run --with python-docx --with svglib --with reportlab --with pillow \
        python docs/training-manual/build_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.graphics import renderPM
from svglib.svglib import svg2rlg

HERE = Path(__file__).parent
FIG = HERE / "figures"
OUT = HERE / "Agon-Practitioners-Manual.docx"

# Palette (general style guide + the Teal-Blue heading override).
CHARCOAL = RGBColor(0x2D, 0x34, 0x36)
TEAL = RGBColor(0x0F, 0x47, 0x61)
STEEL = RGBColor(0x2C, 0x3E, 0x6B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MEDGREY = RGBColor(0x7B, 0x87, 0x94)
WARMGREY_HEX = "F5F5F0"
CODESHADE_HEX = "F1F1EC"

# Glyphs that reportlab's fallback font can't draw — swapped to ASCII for the raster copies only
# (the committed SVGs keep the nicer glyphs for any HTML/markdown rendering).
GLYPH_FIX = {"▶": ">", "▸": ">", "→": "->", "←": "<-",
             "↑": "^", "↓": "v"}

PARTS = [
    ("Agon-Practitioners-Manual-Part-I.md", "## How to use this manual", "## End of Part I"),
    ("Agon-Practitioners-Manual-Part-II.md", "# PART II", "## End of Part II"),
    ("Agon-Practitioners-Manual-Part-III.md", "# PART III", "## End of Part III"),
    ("Agon-Practitioners-Manual-Part-IV.md", "# PART IV", "## End of Part IV"),
    ("Agon-Practitioners-Manual-Part-V.md", "# PART V", "## End of Part V"),
    ("Agon-Practitioners-Manual-Part-VI-and-BackMatter.md", "# PART VI",
     "## End of the manuscript"),
]


# --------------------------------------------------------------------------- figures
def render_figures() -> dict[str, Path]:
    pngs: dict[str, Path] = {}
    for svg in sorted(FIG.glob("*.svg")):
        text = svg.read_text(encoding="utf-8")
        for bad, good in GLYPH_FIX.items():
            text = text.replace(bad, good)
        tmp = svg.with_name(svg.stem + ".raster.svg")
        tmp.write_text(text, encoding="utf-8")
        drawing = svg2rlg(str(tmp))
        scale = 2.0
        drawing.width *= scale
        drawing.height *= scale
        drawing.scale(scale, scale)
        png = svg.with_suffix(".png")
        renderPM.drawToFile(drawing, str(png), fmt="PNG", bg=0xFFFFFF)
        tmp.unlink()
        pngs[svg.name] = png
    return pngs


# --------------------------------------------------------------------------- low-level docx
def shade(el, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    el.append(shd)


def add_runs(p, text: str) -> None:
    """Render inline **bold**, *italic*, and `code` into a paragraph."""
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("`"):
            r = p.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            shade(r._element.get_or_add_rPr(), CODESHADE_HEX)
        elif tok.startswith("**"):
            p.add_run(tok[2:-2]).bold = True
        else:
            p.add_run(tok[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)
    normal.font.color.rgb = CHARCOAL
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.18
    for lvl, sz in {1: 18, 2: 14, 3: 12, 4: 11}.items():
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = "Segoe UI"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = TEAL
        st.paragraph_format.space_before = Pt(14 if lvl <= 2 else 10)
        st.paragraph_format.space_after = Pt(6)
    code = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code.font.size = Pt(9.5)
    code.font.color.rgb = CHARCOAL
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.line_spacing = 1.0
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    header, *body = rows
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, cell in enumerate(header):
        c = t.rows[0].cells[j]
        shade(c._tc.get_or_add_tcPr(), "2C3E6B")
        para = c.paragraphs[0]
        run = para.add_run(re.sub(r"[*`]", "", cell).strip())
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
    for ri, row in enumerate(body):
        cells = t.add_row().cells
        for j, cell in enumerate(row):
            if j >= len(cells):
                break
            para = cells[j].paragraphs[0]
            para.style = doc.styles["Normal"]
            for r in para.runs:
                r.text = ""
            add_runs(para, cell.strip())
            for r in para.runs:
                r.font.size = Pt(10)
            if ri % 2 == 1:
                shade(cells[j]._tc.get_or_add_tcPr(), WARMGREY_HEX)
    doc.add_paragraph()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    note = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Right-click and choose “Update Field” to build the table of contents."
    note.append(t)
    fld.append(note)
    run._element.addprevious(fld)


# --------------------------------------------------------------------------- assembly
def extract(path: Path, start: str, end: str) -> list[str]:
    lines = path.read_text(encoding="utf-8").splitlines()
    si = next(i for i, ln in enumerate(lines) if ln.strip().startswith(start))
    ei = next((i for i, ln in enumerate(lines) if ln.strip().startswith(end)), len(lines))
    return lines[si:ei]


def manuscript_lines() -> list[str]:
    out: list[str] = []
    for fname, start, end in PARTS:
        out += extract(HERE / fname, start, end)
        out.append("")
    return out


def title_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("The Agon Eval Harness")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = TEAL
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A Practitioner's Manual")
    r.font.size = Pt(18)
    r.font.color.rgb = STEEL
    for _ in range(8):
        doc.add_paragraph()
    for label, value in [
        ("Document code", "AGON-TM-001"),
        ("Author", "Samuel R. Taylor"),
        ("Version", "1.0"),
        ("Date", "2026-06-08"),
    ]:
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = line.add_run(f"{label}:  ")
        rl.bold = True
        rl.font.color.rgb = MEDGREY
        line.add_run(value).font.color.rgb = CHARCOAL
    doc.add_page_break()
    h = doc.add_paragraph()
    h.add_run("Contents").bold = True
    h.runs[0].font.size = Pt(14)
    h.runs[0].font.color.rgb = TEAL
    add_toc(doc)
    doc.add_page_break()


def build() -> None:
    pngs = render_figures()
    doc = Document()
    setup_styles(doc)
    title_page(doc)

    lines = manuscript_lines()
    i = 0
    first_heading = True
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # fenced code block
        if stripped.startswith("```"):
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            for j, code_line in enumerate(block):
                p = doc.add_paragraph(style="CodeBlock")
                shade(p._p.get_or_add_pPr(), CODESHADE_HEX)
                p.add_run(code_line if code_line else " ")
                if j == 0:
                    p.paragraph_format.space_before = Pt(6)
            doc.add_paragraph()
            continue

        # heading
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            page_break = (
                level == 1
                or text.startswith("Chapter ")
                or text.startswith("Appendix ")
                or text == "BACK MATTER"
            )
            if page_break and not first_heading:
                doc.add_page_break()
            first_heading = False
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # image
        m = re.match(r"!\[(.*?)\]\((.+?)\)", stripped)
        if m:
            svg_name = Path(m.group(2)).name
            png = pngs.get(svg_name.replace(".png", ".svg"))
            if png and png.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(png), width=Inches(6.3))
            i += 1
            continue

        # table
        if stripped.startswith("|") and i + 1 < len(lines) and set(
            lines[i + 1].strip()
        ) <= set("|-: "):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if set("".join(cells)) <= set("-: "):
                    i += 1
                    continue
                rows.append(cells)
                i += 1
            if rows:
                add_table(doc, rows)
            continue

        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, stripped.lstrip(">").strip())
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = MEDGREY
            i += 1
            continue

        # bullet / numbered list
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # caption (italic line starting with Figure) — centered
        if stripped.startswith("*Figure") and stripped.endswith("*"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, stripped)
            for r in p.runs:
                r.font.size = Pt(9.5)
                r.font.color.rgb = MEDGREY
            i += 1
            continue

        # ordinary paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(str(OUT))
    print(f"wrote {OUT}  ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build()
