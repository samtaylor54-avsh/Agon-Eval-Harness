"""
build_docx.py — Convert measuring-agentic-systems.md into a styled Word handout.

Run with:
    uv run --with python-docx python docs/methodology/build_docx.py
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

# ── Palette ────────────────────────────────────────────────────────────────────
CHARCOAL    = RGBColor(0x2D, 0x34, 0x36)
STEEL_BLUE  = RGBColor(0x2C, 0x3E, 0x6B)
WARM_GREY   = RGBColor(0xF5, 0xF5, 0xF0)
AMBER       = RGBColor(0xD4, 0x88, 0x0F)
MED_GREY    = RGBColor(0x7B, 0x87, 0x94)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# Hex strings (no #) for XML shading elements
HEX_STEEL  = "2C3E6B"
HEX_WARM   = "F5F5F0"
HEX_WHITE  = "FFFFFF"


# ── XML helpers ────────────────────────────────────────────────────────────────

def _set_para_shading(para, fill_hex: str):
    """Apply background shading to a paragraph via w:shd XML."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _set_cell_shading(cell, fill_hex: str):
    """Apply background shading to a table cell via w:shd XML."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


# ── Inline parser ──────────────────────────────────────────────────────────────
# Tokenise **bold**, *italic*, `code`, and plain text spans.
_INLINE_RE = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`)")


def add_inline_runs(para, text: str, base_size: int = 11,
                    base_color: RGBColor = CHARCOAL,
                    base_font: str = "Calibri"):
    """
    Parse *italic*, **bold**, and `code` spans in text and add runs to para.
    Falls back to plain runs for everything else.
    """
    pos = 0
    for m in _INLINE_RE.finditer(text):
        start, end = m.span()
        # Plain text before match
        if pos < start:
            run = para.add_run(text[pos:start])
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color

        token = m.group(0)
        if token.startswith("**"):          # bold
            run = para.add_run(m.group(2))
            run.bold = True
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color
        elif token.startswith("`"):         # inline code
            run = para.add_run(m.group(4))
            run.font.name = "Consolas"
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color
        else:                               # *italic*
            run = para.add_run(m.group(3))
            run.italic = True
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color

        pos = end

    # Trailing plain text
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.name = base_font
        run.font.size = Pt(base_size)
        run.font.color.rgb = base_color


# ── Paragraph helpers ──────────────────────────────────────────────────────────

def _set_line_spacing(para, multiple: float = 1.15):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = multiple


def add_title(doc, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = STEEL_BLUE
    para.paragraph_format.space_after = Pt(4)


def add_subtitle(doc, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = MED_GREY
    _set_line_spacing(para)


def add_standfirst(doc, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = CHARCOAL
    _set_line_spacing(para)
    para.paragraph_format.space_after = Pt(8)


def add_h2(doc, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = CHARCOAL
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)


def add_body(doc, text: str):
    para = doc.add_paragraph()
    add_inline_runs(para, text)
    _set_line_spacing(para)
    para.paragraph_format.space_after = Pt(6)


def add_code_block(doc, lines: list[str]):
    """Each code line gets its own paragraph with Warm Grey shading."""
    # Render as a single paragraph with manual line breaks to avoid excessive spacing
    para = doc.add_paragraph()
    _set_para_shading(para, HEX_WARM)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    for i, line in enumerate(lines):
        if i > 0:
            para.add_run().add_break()  # soft line break inside paragraph
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = CHARCOAL


def add_table(doc, headers: list[str], rows: list[list[str]]):
    """Build the styled pipe table."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for j, hdr in enumerate(headers):
        cell = hdr_row.cells[j]
        _set_cell_shading(cell, HEX_STEEL)
        para = cell.paragraphs[0]
        run = para.add_run(hdr.strip())
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = WHITE

    # Body rows — alternating white / warm grey
    for i, row_data in enumerate(rows):
        fill = HEX_WHITE if i % 2 == 0 else HEX_WARM
        tr = table.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            cell = tr.cells[j]
            _set_cell_shading(cell, fill)
            para = cell.paragraphs[0]
            add_inline_runs(para, cell_text.strip(), base_size=10)


# ── Markdown parser ────────────────────────────────────────────────────────────

def _strip_italic_markers(text: str) -> str:
    """Remove leading/trailing * from an italic-only line."""
    return text.strip().strip("*").strip()


def _parse_table_row(line: str) -> list[str]:
    """Split a `| a | b | c |` line into cells."""
    parts = line.strip().strip("|").split("|")
    return [p.strip() for p in parts]


def _is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r"-+", c) for c in cells)


def parse_and_build(md_path: Path, doc: Document):
    lines = md_path.read_text(encoding="utf-8").splitlines()

    title_done = False
    leading_italics: list[str] = []   # collect up to 2 italic-only lines after H1
    in_code = False
    code_lines: list[str] = []
    table_headers: list[str] | None = None
    table_rows: list[list[str]] = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Fenced code block ─────────────────────────────────────────────────
        if line.strip() == "```":
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                if code_lines:
                    add_code_block(doc, code_lines)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── H1 ────────────────────────────────────────────────────────────────
        if line.startswith("# ") and not title_done:
            add_title(doc, line[2:].strip())
            title_done = True
            i += 1
            continue

        # ── H2 ────────────────────────────────────────────────────────────────
        if line.startswith("## "):
            # Flush any pending table
            if table_headers is not None:
                add_table(doc, table_headers, table_rows)
                table_headers = None
                table_rows = []
            add_h2(doc, line[3:].strip())
            i += 1
            continue

        # ── H3 (defensive) ───────────────────────────────────────────────────
        if line.startswith("### "):
            add_h2(doc, line[4:].strip())  # treat same as H2
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Pipe table ────────────────────────────────────────────────────────
        if line.startswith("|"):
            cells = _parse_table_row(line)
            if table_headers is None:
                table_headers = cells
            elif _is_separator_row(cells):
                pass  # skip separator
            else:
                table_rows.append(cells)
            i += 1
            continue

        # ── Flush completed table when we exit table context ──────────────────
        if table_headers is not None:
            add_table(doc, table_headers, table_rows)
            table_headers = None
            table_rows = []

        # ── Leading italic lines (subtitle / standfirst) ─────────────────────
        if (title_done and len(leading_italics) < 2
                and re.fullmatch(r"\*[^*].+[^*]\*", line.strip())):
            clean = _strip_italic_markers(line)
            if len(leading_italics) == 0:
                add_subtitle(doc, clean)
            else:
                add_standfirst(doc, clean)
            leading_italics.append(clean)
            i += 1
            continue

        # ── Body paragraph ────────────────────────────────────────────────────
        add_body(doc, line.strip())
        i += 1

    # Flush any table that ended at EOF
    if table_headers is not None:
        add_table(doc, table_headers, table_rows)


# ── Normal style baseline ─────────────────────────────────────────────────────

def configure_normal_style(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = CHARCOAL
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    repo_root = Path(__file__).resolve().parent.parent.parent
    md_path   = repo_root / "docs" / "methodology" / "measuring-agentic-systems.md"
    out_path  = repo_root / "docs" / "Agon_Methodology_Essay.docx"

    doc = Document()
    configure_normal_style(doc)

    parse_and_build(md_path, doc)

    doc.save(out_path)
    print(f"Written: {out_path.resolve()}")
